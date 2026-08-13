#!/usr/bin/env python3
"""
Generate the App Manager Deployment Guide & User Manual as a professional DOCX.
Uses python-docx with a Modern Corporate aesthetic (Aptos/Calibri, dark navy
headings, 1.15 line spacing, clean tables, bottom-right page numbers).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===========================================================================
# Style constants -- Modern Corporate aesthetic
# ===========================================================================

# Colors (RGB hex, no #)
COLOR_DARK_NAVY = "1F3864"
COLOR_BODY      = "333333"
COLOR_GRAY_TEXT = "595959"
COLOR_PAGE_NUM  = "808080"

# Fonts
FONT_HEADING = "Aptos Display"
FONT_BODY    = "Calibri"

# Font sizes (points)
SZ_BODY      = 11
SZ_H1        = 20
SZ_H2        = 16
SZ_H3        = 13
SZ_H4        = 11
SZ_CAPTION   = 9
SZ_PAGE_NUM  = 9

# Spacing (Pt)
SPACE_AFTER_BODY   = 8
SPACE_BEFORE_H1    = 24
SPACE_BEFORE_H2    = 18
SPACE_BEFORE_H3    = 12


# ===========================================================================
# Core helpers
# ===========================================================================

def create_character_style(doc, name, font_name, size_pt, color_hex, bold=False):
    from docx.enum.style import WD_STYLE_TYPE
    style = doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor.from_string(color_hex)
    style.font.bold = bold


def add_body_text(doc, paragraphs):
    if isinstance(paragraphs, str):
        paragraphs = [paragraphs]
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.style = "Normal"


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")


def add_step_list(doc, steps):
    """Numbered step list: (label, description) tuples."""
    for label, desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        run_label = p.add_run(label)
        run_label.font.bold = True
        run_label.font.name = FONT_BODY
        run_label.font.size = Pt(SZ_BODY)
        run_label.font.color.rgb = RGBColor.from_string(COLOR_DARK_NAVY)
        p.add_run(" " + desc)


def add_code_block(doc, code_text):
    """Code block via 1-cell shaded table (reliable background)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    col = table.columns[0]
    col.width = Twips(9000)

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F8F9FA')
    tcPr.append(shd)

    tcBorders = OxmlElement('w:tblCellMar')
    for edge in ['top', 'left', 'bottom', 'right']:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:w'), '120')
        elem.set(qn('w:type'), 'dxa')
        tcBorders.append(elem)
    tcPr.append(tcBorders)

    tblPr = table._tbl.tblPr  # auto-creates if missing
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), 'D0D7DE')
        tblBorders.append(elem)
    tblPr.append(tblBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("333333")


def add_note_box(doc, title, text):
    """Callout box: indented, italic, with colored text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    run = p.add_run("\u26a0\ufe0f " + title + ": " + text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("0F4C81")
    run.italic = True


def _add_field_code(run, field_code, result_text=""):
    """Inject a Word field code (PAGE, NUMPAGES, TOC) via raw XML."""
    r = run._r
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    r.append(fc_begin)
    instr = OxmlElement('w:instrText')
    instr.text = field_code
    r.append(instr)
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r.append(fc_sep)
    if result_text:
        t = OxmlElement('w:t')
        t.text = result_text
        r.append(t)
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r.append(fc_end)


def add_styled_table(doc, headers, rows, alignments=None):
    """Add a professional table with header row and optional column alignment."""
    cols = len(headers)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].alignment = (
            alignments[i] if alignments else WD_PARAGRAPH_ALIGNMENT.LEFT
        )
        # Bold header
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True

    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)

    return table


# ===========================================================================
# Document creation
# ===========================================================================

def create_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin  = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    setup_styles(doc)
    build_cover_page(doc)
    build_toc(doc)
    build_deployment_guide(doc)
    build_user_manual(doc)
    build_api_reference(doc)
    build_db_schema(doc)
    setup_footer(doc, section)

    output_path = "App_Manager_Deployment_Guide_and_User_Manual.docx"
    doc.save(output_path)
    abs_path = os.path.abspath(output_path)
    print(f"Document saved: {abs_path}")
    print(f"Size: {os.path.getsize(output_path) / 1024:.1f} KB")
    return abs_path


# ===========================================================================
# Styles
# ===========================================================================

def setup_styles(doc):
    styles = doc.styles

    # ── Normal (body default) ──
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(SZ_BODY)
    normal.font.color.rgb = RGBColor.from_string(COLOR_BODY)
    p = normal.paragraph_format
    p.line_spacing = 1.15
    p.space_after = Pt(SPACE_AFTER_BODY)

    # ── Heading 1 ──
    h1 = styles["Heading 1"]
    h1.font.name = FONT_HEADING
    h1.font.size = Pt(SZ_H1)
    h1.font.color.rgb = RGBColor.from_string(COLOR_DARK_NAVY)
    h1.font.bold = False
    pf = h1.paragraph_format
    pf.space_before = Pt(SPACE_BEFORE_H1)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    pf.keep_together = True
    pf.page_break_before = True

    # ── Heading 2 ──
    h2 = styles["Heading 2"]
    h2.font.name = FONT_HEADING
    h2.font.size = Pt(SZ_H2)
    h2.font.color.rgb = RGBColor.from_string(COLOR_DARK_NAVY)
    h2.font.bold = False
    pf = h2.paragraph_format
    pf.space_before = Pt(SPACE_BEFORE_H2)
    pf.space_after = Pt(6)
    pf.keep_with_next = True

    # ── Heading 3 ──
    h3 = styles["Heading 3"]
    h3.font.name = FONT_HEADING
    h3.font.size = Pt(SZ_H3)
    h3.font.color.rgb = RGBColor.from_string(COLOR_DARK_NAVY)
    h3.font.bold = True
    pf = h3.paragraph_format
    pf.space_before = Pt(SPACE_BEFORE_H3)
    pf.space_after = Pt(4)
    pf.keep_with_next = True

    # ── Heading 4 ──
    h4 = styles["Heading 4"]
    h4.font.name = FONT_BODY
    h4.font.size = Pt(SZ_H4)
    h4.font.color.rgb = RGBColor.from_string(COLOR_DARK_NAVY)
    h4.font.bold = True
    pf = h4.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    pf.keep_with_next = True

    # ── Caption ──
    cap = styles["Caption"]
    cap.font.name = FONT_BODY
    cap.font.size = Pt(SZ_CAPTION)
    cap.font.color.rgb = RGBColor.from_string(COLOR_GRAY_TEXT)
    cap.font.italic = True
    pf = cap.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Character styles
    create_character_style(doc, "CodeInline", "Consolas", 9, "4A90D9")


# ===========================================================================
# Cover page
# ===========================================================================

def build_cover_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("App Manager")
    run.font.name = FONT_HEADING
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLOR_DARK_NAVY)
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = subtitle.add_run("Panduan Penerapan & Manual Pengguna")
    run2.font.name = FONT_BODY
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor.from_string(COLOR_BODY)
    subtitle.paragraph_format.space_after = Pt(36)

    doc.add_paragraph()

    info_table = doc.add_table(rows=0, cols=1)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = True

    info_items = [
        ("Projek", "App Register Manager (app-manager)"),
        ("Persekitaran", "Laragon, Windows, Apache + PHP 7.4 + MySQL 8"),
        ("Frontend", "Vue 3.5 + Vue Router 4, Vite 6.2"),
        ("Backend", "PHP REST API (api/apps.php, api/scan.php)"),
        ("Pangkalan Data", "MySQL 8, pangkalan data app_manager"),
        ("Pengguna", "Mohd Ilhammuddin Mohd Fuead"),
        ("Tarikh", "Ogos 2026"),
    ]

    for label, value in info_items:
        cell = info_table.add_row().cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.font.name = FONT_BODY
        r1.font.size = Pt(11)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor.from_string(COLOR_DARK_NAVY)
        r2 = p.add_run(value)
        r2.font.name = FONT_BODY
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor.from_string(COLOR_BODY)

    doc.add_paragraph()
    doc.add_paragraph()

    note = doc.add_paragraph()
    note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nr = note.add_run(
        "Dokumen ini mencakup semua maklumat yang diperlukan untuk "
        "memasang, mengguna, dan mentadbir App Manager dari permulaan."
    )
    nr.font.name = FONT_BODY
    nr.font.size = Pt(10)
    nr.font.color.rgb = RGBColor.from_string("7F8C89")
    note.paragraph_format.space_before = Pt(48)

    doc.add_page_break()


# ===========================================================================
# Table of Contents
# ===========================================================================

def build_toc(doc):
    """Insert a Word TOC field that auto-generates from heading styles."""
    doc.add_paragraph("Isi Kandungan", style="Heading 1")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)

    run = p.add_run()
    r = run._r

    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    r.append(fc_begin)

    instr = OxmlElement('w:instrText')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r.append(instr)

    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r.append(fc_sep)

    placeholder = OxmlElement('w:t')
    placeholder.text = 'Tekan Ctrl+A, kemudian F9 untuk menggupdate daftar isi.'
    r.append(placeholder)

    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r.append(fc_end)


# ===========================================================================
# Section 1: Deployment Guide
# ===========================================================================

def build_deployment_guide(doc):
    doc.add_paragraph("1. Panduan Penerapan (Dari Permulaan)", style="Heading 1")

    add_body_text(doc, [
        "Panduan ini memandu anda melalui proses lengkap untuk memasang "
        "App Manager dari permulaan — dari prasyarat sistem sehingga akses penuh "
        "di pelayar tempatan."
    ])

    # 1.1 Prerequisites
    doc.add_paragraph("1.1 Prasyarat Sistem", style="Heading 2")
    add_body_text(doc, [
        "Sistem berikut diperlukan untuk menjalankan App Manager:"
    ])
    add_bullet_list(doc, [
        "Laragon (termasuk Apache, PHP 7.4, dan MySQL 8) — paket penuh",
        "Node.js 18.x (untuk membina frontend Vue 3)",
        "Git (untuk pengeskakan SCM projek)",
        "Windows dengan akses ke port 80 (Apache)",
        "Ruang cak RAM — pastikan folder C:\\laragon\\www\\ boleh ditulis",
    ])
    add_note_box(doc, "Nota Teknikal",
        "Pangkalan data app_manager tidak memerlukan password untuk "
        "pengguna root pada konfigurasi Laragon lalai. Semua sambungan PDO "
        "menggunakan UTF-8 (utf8mb4) untuk sokongan penuh aksara Melayu.")

    # 1.2 Project Setup
    doc.add_paragraph("1.2 Persediaan Projek", style="Heading 2")
    add_body_text(doc, [
        "Langkah pertama ialah menyediakan folder projek dan memasang pergantungan."
    ])
    add_step_list(doc, [
        ("1.", "Cipta folder projek di bawah Laragon: C:\\laragon\\www\\app-manager"),
        ("2.", "Salin atau klon kod sumber ke dalam folder"),
        ("3.", "Pasang pergantungan Node.js dengan npm install"),
        ("4.", "Sediakan dan import pangkalan data MySQL"),
        ("5.", "Konfigurasikan Apache virtual host"),
    ])
    add_code_block(doc,
        "cd C:\\laragon\\www\\app-manager\n"
        "git clone <repo-url> .\n"
        "npm install")

    add_body_text(doc, [
        "Struktur folder hasil:"
    ])
    add_code_block(doc,
        "app-manager/\n"
        "├── api/\n"
        "│   ├── apps.php              # REST API — CRUD aplikasi\n"
        "│   ├── scan.php              # REST API — pindai direktori Git\n"
        "│   ├── schema.sql            # Definisi penuh skema pangkalan data\n"
        "│   ├── migrate_active_apps.sql # Migrasi kolom is_active\n"
        "│   └── migrate_pin_apps.sql    # Migrasi kolom is_pinned\n"
        "├── src/\n"
        "│   ├── main.js               # Titik masuk Vue 3 SPA\n"
        "│   ├── App.vue               # Komponen induk\n"
        "│   ├── api/index.js          # Klien API (fetch wrapper)\n"
        "│   ├── config.js             # URL builder (editor/pelayar)\n"
        "│   ├── router/index.js       # Vue Router (hash mode)\n"
        "│   └── views/\n"
        "│       ├── Dashboard.vue\n"
        "│       ├── Detail.vue\n"
        "│       └── AddApp.vue\n"
        "├── dist/                     # Hasil binaan production\n"
        "├── styles/main.css           # Gaya Morandi (palet yang sopan)\n"
        "├── config.php                # Konfigurasi pangkalan data\n"
        "├── vite.config.js            # Konfigurasi Vite\n"
        "├── index.html                # Entry point\n"
        "└── package.json              # Manifest & pergantungan")

    # 1.3 Database
    doc.add_paragraph("1.3 Penyediaan Pangkalan Data", style="Heading 2")
    add_body_text(doc, [
        "App Manager menggunakan satu pangkalan data MySQL bernama app_manager "
        "dengan 11 jadual. Skema lengkap terletak di api/schema.sql. "
        "Untuk pangkalan data baru, jalankan skema penuh; untuk pangkalan data "
        "sedia ada, jalankan migrasi."
    ])
    add_code_block(doc,
        "# Buat / reset pangkalan data dari skema\n"
        "mysql -u root < api/schema.sql\n"
        "\n"
        "# Untuk pangkalan data sedia ada (upgrade):\n"
        "mysql -u root app_manager -e \"$(cat api/migrate_active_apps.sql)\"\n"
        "mysql -u root app_manager -e \"$(cat api/migrate_pin_apps.sql)\"")

    add_body_text(doc, [
        "Maklumat penting tentang migrasi:"
    ])
    add_bullet_list(doc, [
        "is_active — kolom TINYINT(1) DEFAULT 1. Semua aplikasi baru secara lalai aktif.",
        "is_pinned — kolom TINYINT(1) DEFAULT 0. Semua aplikasi tidak disematkan secara lalai.",
        "Indeks idx_apps_active dan idx_apps_pinned ditambah untuk prestasi yang lancar.",
        "Migrasi boleh dijalankan berulang kali — ia memeriksa information_schema sebelum ALTER.",
    ])
    add_note_box(doc, "PowerShell",
        "PowerShell tidak menyokong sintaks mysql < file.sql secara langsung. "
        "Guna:\n"
        "  Get-Content api/migrate_active_apps.sql -Raw | mysql -u root app_manager\n"
        "atau\n"
        "  & mysql -u root app_manager -e (Get-Content api/migrate_active_apps.sql -Raw)")

    # 1.4 API Configuration
    doc.add_paragraph("1.4 Konfigurasi API (config.php)", style="Heading 2")
    add_body_text(doc, [
        "Fail config.php di akar projek mengkonfigurasi sambungan pangkalan data. "
        "Untuk persekitaran Laragon, nilai lalai biasanya cukup. Ubah jika perlu."
    ])
    add_code_block(doc,
        "<?php\n"
        "// config.php\n"
        "define('DB_HOST', '127.0.0.1');\n"
        "define('DB_PORT', '3306');\n"
        "define('DB_NAME', 'app_manager');\n"
        "define('DB_USER', 'root');\n"
        "define('DB_PASS', '');")

    add_body_text(doc, [
        "Uji sambungan API:"
    ])
    add_code_block(doc, "curl http://localhost/api/apps.php")

    # 1.5 Frontend
    doc.add_paragraph("1.5 Persediaan Frontend", style="Heading 2")
    add_body_text(doc, [
        "App Manager adalah aplikasi Vue 3 single-page application yang dibina "
        "menggunakan Vite. Untuk persekitaran pembangunan, gunakan pelayan dev Vite "
        "yang menyalurkan permintaan API ke Apache secara automatik."
    ])
    add_step_list(doc, [
        ("1.", "Pasang pergantungan dengan npm install"),
        ("2.", "Jalankan pelayan pembangunan dengan npm run dev"),
        ("3.", "Akses aplikasi di http://localhost:5173/app-manager/dist/"),
    ])
    add_code_block(doc,
        "npm install             # Pasang Vue 3, Vue Router, Vite\n"
        "npm run dev             # Pelayan dev @ http://localhost:5173/app-manager/dist/\n"
        "npm run build           # Bina untuk production -> dist/")

    add_body_text(doc, [
        "Konfigurasi Vite (vite.config.js) menggunakan base path /app-manager/dist/ "
        "dan menyalurkan permintaan API /app-manager/api/* ke http://localhost (Apache)."
    ])

    # 1.6 Apache
    doc.add_paragraph("1.6 Konfigurasi Apache Virtual Host", style="Heading 2")
    add_body_text(doc, [
        "Aplikasi diakses melalui http://localhost/app-manager/ (bukan domain .test — "
        "McAfee Web Gateway menyekat resolusi TLD .test). Folder projek menjadi folder "
        "simpanan Apache yang automatik."
    ])
    add_code_block(doc,
        "# Pastikan struktur folder:\n"
        "# C:\\laragon\\www\\app-manager\\\n"
        "#   index.html  (dari npm run build)\n"
        "#   dist/       (aset CSS/JS)\n"
        "#   api/        (PHP API)\n"
        "#\n"
        "# Apache menyajikan:\n"
        "#   http://localhost/app-manager/       -> index.html\n"
        "#   http://localhost/api/apps.php       -> api/apps.php\n"
        "#   http://localhost/app-manager/dist/  -> dist/assets/")

    # 1.7 Production Build
    doc.add_paragraph("1.7 Pembinaan untuk Produksi", style="Heading 2")
    add_body_text(doc, [
        "Setelah semua pergantungan dipasang dan pangkalan data disediakan, "
        "bina aplikasi untuk penggunaan produksi."
    ])
    add_code_block(doc,
        "npm run build\n"
        "# Output:\n"
        "#   dist/index.html\n"
        "#   dist/assets/index-CaypCsXY.js  (121.6 kB)\n"
        "#   dist/assets/index-vpmH94NN.css (13.3 kB)")

    add_body_text(doc, [
        "Fail dist/index.html ialah titik masuk. Aset CSS/JS dirujuk "
        "menggunakan jejak relatif /app-manager/dist/assets/ "
        "(ditentukan oleh nilai base dalam vite.config.js)."
    ])

    add_note_box(doc, "Pengesahan Akhir",
        "Setelah semua langkah selesai, akses http://localhost/app-manager/. "
        "Dashboard harus memaparkan kad-kad aplikasi. Jika kad tidak kelihatan, pastikan:\n"
        "1. MySQL berjalan (Laragon daemon)\n"
        "2. Pangkalan data app_manager wujud\n"
        "3. Skema telah diimport\n"
        "4. Vite dev server atau Apache dist/ sedang berfungsi")

    # 1.8 API Quick Reference
    doc.add_paragraph("1.8 Rujukan Pantas API", style="Heading 2")
    add_body_text(doc, [
        "API tersedia di dua endpoint — lihat Bahagian 3 untuk maklumat terperinci."
    ])
    add_code_block(doc,
        "GET  /api/apps.php                     — senarai semua aplikasi\n"
        "GET  /api/apps.php?name=X               — maklumat terperinci satu aplikasi\n"
        "POST /api/apps.php                     — tambah aplikasi (body: name, path)\n"
        "PUT  /api/apps.php?name=X               — kemaskini profil (body: active, pinned, etc.)\n"
        "DELETE /api/apps.php?name=X&note_id=Y   — padam catatan\n"
        "GET  /api/scan.php                      — pindai semua projek\n"
        "GET  /api/scan.php?path=D               — pindai satu projek sahaja")


# ===========================================================================
# Section 2: User Manual
# ===========================================================================

def build_user_manual(doc):
    doc.add_paragraph("2. Manual Pengguna", style="Heading 1")
    add_body_text(doc, [
        "Manual ini memandu anda melalui semua ciri App Manager — dari melihat "
        "dashboard sehingga ke pengurusan catatan dan perkhidmatan. "
        "Antarmuka pengguna berada dalam Bahasa Melayu."
    ])

    # 2.1 Dashboard
    doc.add_paragraph("2.1 Papan Pemandangan (Dashboard)", style="Heading 2")
    add_body_text(doc, [
        "Dashboard ialah skrin utama. Strukturnya terdiri daripada:"
    ])
    add_bullet_list(doc, [
        "Bar header — butang Segar Semula (refresh) dan Tambah Aplikasi",
        "Bar statistik — jumlah aplikasi, bersih/kotor, masa pemindaian terakhir",
        "Bar carian & penapis — cari aplikasi, atau tapis mengikut stack (PHP, Node, Vue) atau status SCM",
        "Grid kad aplikasi — bahagikan kepada bahagian Aktif dan Tidak Aktif",
    ])
    add_code_block(doc, "URL: http://localhost/app-manager/")

    # 2.2 App Cards
    doc.add_paragraph("2.2 Kad Aplikasi", style="Heading 2")
    add_body_text(doc, [
        "Setiap kad mewakili satu projek di C:\\laragon\\www\\. "
        "Kad menampilkan maklumat ringkas:"
    ])
    add_bullet_list(doc, [
        "Nama aplikasi (dengan ikon pin jika disematkan)",
        "Badge stack — PHP, Laravel, Vue, Node, dll.",
        "Maklumat SCM — cawangan Git, masa komit terakhir, URL remote",
        "Pratonton catatan (jika ada)",
        "Butang aksi — buka di editor, buka di pelayar",
    ])

    # 2.3 Active/Inactive
    doc.add_paragraph("2.3 Status Aktif / Tidak Aktif", style="Heading 2")
    add_body_text(doc, [
        "Setiap aplikasi mempunyai status aktif atau tidak aktif. Status ini "
        "mengawal sama ada kad kelihatan di bahagian atas dashboard "
        "(aktif) atau di bawah pem_separator 'APLIKASI TIDAK AKTIF' "
        "(tidak aktif). Kad yang tidak aktif juga akan redup (opacity 0.55)."
    ])

    add_code_block(doc,
        "+---------------- Dashboard ----------------+\n"
        "|  [kad: laravel-api]  [kad: vue-app]     |\n"
        "|  [kad: nextjs-app]   [kad: ecaller-dev] |\n"
        "|  ... 19 aplikasi aktif ...               |\n"
        "|                                        |\n"
        "|  --------- APLIKASI TIDAK AKTIF --------|\n"
        "|                                        |\n"
        "|  [kad: app-manager]  (redup, opacity 0.55) |\n"
        "|  [kad: novel-manager] (redup)             |\n"
        "+------------------------------------------+")

    add_body_text(doc, [
        "Cara mengubah status:"
    ])
    add_step_list(doc, [
        ("1.", "Klik butang mata (icon mata / mata bercelah) pada mana-mana kad"),
        ("2.", "Status berubah serta-merta (optimistic update di dashboard)"),
        ("3.", "Permintaan PUT dihantar ke API: { \"active\": false }"),
        ("4.", "Vue reaktif memindahkan kad ke bahagian yang betul"),
    ])
    add_note_box(doc, "Ketahanan Status",
        "Status aktif/tidak aktif TIDAK dikemaskini oleh pindai (scan.php). "
        "Ini bermaksud pilihan anda kekal selepas Segar Semula atau Imbas Semula. "
        "Hanya maklumat SCM (branch, commit, changed files) yang segar semula.")

    # 2.4 Pin
    doc.add_paragraph("2.4 Menyematkan Aplikasi (Pin)", style="Heading 2")
    add_body_text(doc, [
        "Menyematkan aplikasi menempatkannya sentiasa di bahagian atas dalam "
        "kumpulan aktif atau tidak aktif. Ini berguna untuk menyimpan akses "
        "pantas ke projek yang selalu digunakan."
    ])
    add_code_block(doc,
        "Dashboard -> kad aplikasi -> klik ikon pin (icon pin)\n"
        "PUT /api/apps.php?name=<nama> -> { \"pinned\": true }")

    # 2.5 Search & Filter
    doc.add_paragraph("2.5 Carian & Penapis", style="Heading 2")
    add_body_text(doc, [
        "Bar carian di bawah header membolehkan menapis aplikasi."
    ])
    add_bullet_list(doc, [
        "Teks — cari perkata kunci dalam nama aplikasi, framework, atau cawangan Git",
        "Stack — tapis mengikut jenis (PHP, Node, Vue)",
        "Status SCM — tapis mengikut bersih (clean) atau kotor (dirty)",
    ])
    add_code_block(doc,
        "Contoh carian: ketik \"laravel\" -> menampilkan app-manager, report-api, retail-app...")

    # 2.6 Detail
    doc.add_paragraph("2.6 Halaman Terperinci (Detail)", style="Heading 2")
    add_body_text(doc, [
        "Klik mana-mana kad untuk membuka halaman terperinci. "
        "Halaman ini menampilkan maklumat penuh aplikasi."
    ])
    add_step_list(doc, [
        ("1.", "Butang kembali — kembali ke dashboard"),
        ("2.", "Nama aplikasi + butang mata aktif/tidak aktif"),
        ("3.", "Badge stack — versi bahasa dan framework"),
        ("4.", "Panel kiri: Stack (composer.json / package.json)"),
        ("5.", "Panel kanan: SCM, Perkhidmatan, Catatan"),
        ("6.", "Butang aksi: Buka di Editor, Buka di Pelayar, Imbas Semula"),
    ])
    add_body_text(doc, [
        "Pada halaman terperinci, perubahan status aktif/tidak aktif menyegarkan "
        "objek aplikasi penuh dari API (bukan optimistic update — ini adalah "
        "pandangan kebenaran sumber)."
    ])

    # 2.7 Notes
    doc.add_paragraph("2.7 Catatan & Ulasan (Notes)", style="Heading 2")
    add_body_text(doc, [
        "Setiap aplikasi boleh mempunyai banyak catatan — ini ialah jurnal "
        "tempoh untuk ulasan, perubahan, atau idea. Catatan disimpan di "
        "jadual app_notes dengan stamping masa otomatik."
    ])
    add_step_list(doc, [
        ("1.", "Pergi ke halaman terperinci aplikasi"),
        ("2.", "Skrol ke bahagian 'Catatan & Ulasan'"),
        ("3.", "Taip nota di kotak teks di bawah"),
        ("4.", "Klik 'Tambah Catatan' — tarikh & masa dimasukkan otomatis"),
        ("5.", "Untuk memadam: klik icon sisa (icon sisa) pada mana-mana nota"),
    ])

    # 2.8 Services
    doc.add_paragraph("2.8 Perkhidmatan Pihak Ketiga (Services)", style="Heading 2")
    add_body_text(doc, [
        "Daftar perkhidmatan pihak ketiga yang digunakan aplikasi — "
        "contohnya Firebase (auth), Redis (cache), Stripe (payment), dan lain-lain."
    ])
    add_step_list(doc, [
        ("1.", "Di halaman terperinci, klik 'Edit' pada bahagian Perkhidmatan"),
        ("2.", "Klik 'Tambah' untuk menambah baris perkhidmatan"),
        ("3.", "Isi: Nama, Jenis (dropdown), Pembekal, URL endpoint, Nota"),
        ("4.", "Klik 'Simpan' — perkhidmatan disimpan ke jadual app_services"),
    ])
    add_code_block(doc,
        "Jenis perkhidmatan yang disokong:\n"
        "auth, database, cache, storage, email, payment, sms,\n"
        "api, monitoring, search, queue, cdn, other")

    # 2.9 Scanning
    doc.add_paragraph("2.9 Memindai Projek (Scanning)", style="Heading 2")
    add_body_text(doc, [
        "App Manager secara otomatis mengeskakan teknologi projek (PHP/Node/Python) "
        "daripada composer.json / package.json, dan maklumat SCM daripada .git/. "
        "Terdapat dua cara untuk memindai:"
    ])
    add_bullet_list(doc, [
        "Segar Semula — pindai semua projek di C:\\laragon\\www\\",
        "Imbas Semula — pindai satu projek sahaja (dari halaman terperinci)",
    ])
    add_code_block(doc,
        "# Segar Semula (dari dashboard)\n"
        "GET /api/scan.php\n"
        "\n"
        "# Imbas Semula (dari halaman detail)\n"
        "GET /api/scan.php?path=C:\\laragon\\www\\laravel-api")

    add_body_text(doc, [
        "Proses pemandian:"
    ])
    add_step_list(doc, [
        ("1.", "scan.php mencrawl direktori www/ satu tahap"),
        ("2.", "Untuk setiap folder: baca composer.json, package.json, .git/"),
        ("3.", "Transaksi MySQL: upsert app -> refresh stacks -> upsert scm -> refresh changed_files -> update scan_log"),
        ("4.", "Frontend re-fetch apps.php untuk data segar"),
    ])

    # 2.10 Add App
    doc.add_paragraph("2.10 Menambah Aplikasi Secara Manual", style="Heading 2")
    add_body_text(doc, [
        "Jika projek tiada di dalam folder yang dicsrawl, anda boleh menambahkannya secara manual."
    ])
    add_step_list(doc, [
        ("1.", "Klik 'Tambah Aplikasi' di bar header dashboard"),
        ("2.", "Isi: Nama, Path (folder sebenar), Nota (opsional)"),
        ("3.", "Klik Hantar — aplikasi ditambah dengan is_active=1, is_pinned=0"),
        ("4.", "Jalankan 'Segar Semula' untuk mengeskakan stack dan SCM"),
    ])


# ===========================================================================
# Section 3: API Reference
# ===========================================================================

def build_api_reference(doc):
    doc.add_paragraph("3. Rujukan API", style="Heading 1")
    add_body_text(doc, [
        "App Manager mengekspor REST API sederhana yang menggunakan JSON. "
        "Semua endpoint menggunakan pangkalan data app_manager yang disediakan. "
        "Respons dikembalikan dengan JSON_UNESCAPED_UNICODE untuk sokongan penuh aksara Melayu."
    ])

    # 3.1 apps.php
    doc.add_paragraph("3.1 api/apps.php", style="Heading 2")
    add_body_text(doc, [
        "Endpoint utama — menguruskan rekod aplikasi, catatan, dan status aktif/pin."
    ])

    add_styled_table(
        doc,
        ["Method", "Endpoint", "Deskripsi", "Body", "Respons"],
        [
            ("GET", "/api/apps.php", "Senarai semua aplikasi", "-",
             "JSON: { count, apps[], last_scan_at }"),
            ("GET", "/api/apps.php?name=X", "Satu aplikasi (terperinci)", "-",
             "JSON: app object penuh"),
            ("POST", "/api/apps.php", "Tambah aplikasi baru", "name, path, notes?",
             "201: { message, app }"),
            ("POST", "/api/apps.php?name=X", "Tambah catatan", "note",
             "201: { message, note }"),
            ("PUT", "/api/apps.php?name=X", "Kemaskini profil", "name? path? notes? active? pinned? services?",
             "JSON: { message, app }"),
            ("DELETE", "/api/apps.php?name=X&note_id=Y", "Padam catatan", "-",
             "JSON: { message, deleted_id }"),
        ])

    # 3.2 scan.php
    doc.add_paragraph("3.2 api/scan.php", style="Heading 2")
    add_body_text(doc, [
        "Endpoint pemindaian — mengeskakan teknologi dan SCM untuk projek."
    ])

    add_styled_table(
        doc,
        ["Method", "Endpoint", "Deskripsi", "Respons"],
        [
            ("GET", "/api/scan.php", "Pindai semua projek di C:\\laragon\\www\\",
             "JSON: { scanned, apps[] }"),
            ("GET", "/api/scan.php?path=D", "Pindai satu projek sahaja",
             "JSON: { scanned, app }"),
        ])

    add_note_box(doc, "Contoh Panggilan API",
        "Mengaktifkan semula aplikasi dari tidak aktif:")
    add_code_block(doc,
        "curl -X PUT http://localhost/api/apps.php?name=app-manager \\\n"
        "  -H 'Content-Type: application/json' \\\n"
        "  -d '{\"active\": true}'")

    add_body_text(doc, [
        "Respons:"
    ])
    add_code_block(doc,
        "{\n"
        "  \"message\": \"App profile updated\",\n"
        "  \"app\": {\n"
        "    \"id\": 97,\n"
        "    \"name\": \"app-manager\",\n"
        "    \"is_active\": 1,\n"
        "    \"is_pinned\": 0,\n"
        "    ...\n"
        "  }\n"
        "}")

    # 3.3 Filter parameters
    doc.add_paragraph("3.3 Parameter Penapis (GET apps.php)", style="Heading 2")
    add_body_text(doc, [
        "Senarai aplikasi boleh disaring menggunakan parameter query:"
    ])

    add_styled_table(
        doc,
        ["Parameter", "Nilai", "Deskripsi"],
        [
            ("?search=", "q", "Cari perkata kunci dalam nama, framework, branch"),
            ("?status=", "dirty|clean", "Tapis mengikut status SCM"),
            ("?stack=", "php|node|python|other", "Tapis mengikut jenis stack"),
        ])


# ===========================================================================
# Section 4: Database Schema
# ===========================================================================

def build_db_schema(doc):
    doc.add_paragraph("4. Skema Pangkalan Data", style="Heading 1")
    add_body_text(doc, [
        "Pangkalan data app_manager mengandungi 11 jadual. "
        "Berikut adalah rujukan lengkap yang disusun mengikut imej hubungan."
    ])

    # 4.1 apps table
    doc.add_paragraph("4.1 Jadual apps (Entiti Utama)", style="Heading 2")
    add_body_text(doc, [
        "Jadual induk untuk semua projek. Setiap baris ialah satu aplikasi."
    ])

    add_styled_table(
        doc,
        ["Kolom", "Jenis", "Keterangan"],
        [
            ("id", "INT PK AUTO_INC", "Pengenal unik"),
            ("name", "VARCHAR(120) UNIQUE", "Nama aplikasi, juga pengenal URL"),
            ("path", "VARCHAR(512)", "Laluan penuh ke folder projek"),
            ("notes", "TEXT", "Nota umum tentang projek"),
            ("is_active", "TINYINT(1) DEFAULT 1", "1 = aktif, 0 = tidak aktif"),
            ("is_pinned", "TINYINT(1) DEFAULT 0", "1 = disematkan ke atas dashboard"),
            ("pinned_at", "TIMESTAMP NULL", "Masa penyematian"),
            ("created_at", "TIMESTAMP", "Waktu penciptaan rekod"),
            ("updated_at", "TIMESTAMP", "Auto-update pada perubahan"),
        ],
        [WD_PARAGRAPH_ALIGNMENT.LEFT, WD_PARAGRAPH_ALIGNMENT.LEFT,
         WD_PARAGRAPH_ALIGNMENT.LEFT])

    add_body_text(doc, [
        "Indeks: idx_apps_updated_at, idx_apps_pinned (is_pinned, pinned_at), "
        "idx_apps_active (is_active). Keunikan: uniq_apps_name."
    ])

    # 4.2 ERD
    doc.add_paragraph("4.2 Imej Hubungan (ERD)", style="Heading 2")
    add_body_text(doc, [
        "Jadual apps adalah entiti induk — semua jadual lain merujuk kepadanya melalui foreign key."
    ])
    add_code_block(doc,
        "          apps\n"
        "            | (1)\n"
        "            |\n"
        "    +-------+-------+\n"
        "    |               |\n"
        " (1:M)           (1:1)\n"
        "    |               |\n"
        "    v               v\n"
        " app_stack    app_scm --- (1:M) --> app_changed_files\n"
        "    |\n"
        " (1:M)\n"
        "    |\n"
        "    v\n"
        " app_services\n"
        "    |\n"
        "    |\n"
        " (1:M)\n"
        "    |\n"
        "    v\n"
        " app_notes\n"
        "\n"
        " app_scan_log (1:1 dengan apps)")

    # 4.3 All tables summary
    doc.add_paragraph("4.3 Ringkasan Semua Jadual", style="Heading 2")

    add_styled_table(
        doc,
        ["Jadual", "Kunci", "Keterangan"],
        [
            ("apps", "INT PK, 9 kolom", "Rekod utama setiap projek"),
            ("app_scm", "INT PK + app_id FK", "Maklumat Git: remote, branch, commit, status"),
            ("app_stack", "INT PK + app_id FK", "Teknologi: PHP/Node/Python + framework + dependencies"),
            ("app_changed_files", "INT PK + scm_id FK", "Fail yang berubah — snapshot setiap scan"),
            ("app_services", "INT PK + app_id FK", "Perkhidmatan pihak ketiga — manual entry"),
            ("app_scan_log", "app_id PK FK", "Masa scan terakhir per app"),
            ("app_notes", "INT PK + app_id FK", "Journal catatan/ulasan per app"),
        ])


# ===========================================================================
# Footer
# ===========================================================================

def setup_footer(doc, section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    run_label = p.add_run("Halaman ")
    run_label.font.name = FONT_BODY
    run_label.font.size = Pt(SZ_PAGE_NUM)
    run_label.font.color.rgb = RGBColor.from_string(COLOR_PAGE_NUM)

    run_page = p.add_run()
    run_page.font.name = FONT_BODY
    run_page.font.size = Pt(SZ_PAGE_NUM)
    _add_field_code(run_page, "PAGE", "1")

    run_slash = p.add_run(" / ")
    run_slash.font.name = FONT_BODY
    run_slash.font.size = Pt(SZ_PAGE_NUM)
    run_slash.font.color.rgb = RGBColor.from_string(COLOR_PAGE_NUM)

    run_total = p.add_run()
    run_total.font.name = FONT_BODY
    run_total.font.size = Pt(SZ_PAGE_NUM)
    run_total.font.color.rgb = RGBColor.from_string(COLOR_PAGE_NUM)
    _add_field_code(run_total, "NUMPAGES", "1")


# ===========================================================================
# Main
# ===========================================================================

if __name__ == "__main__":
    output = create_document()
    print(f"\nDocument generated successfully!")
